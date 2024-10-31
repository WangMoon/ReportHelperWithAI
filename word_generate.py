from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.shared import  Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import  qn
from docx.shared import Inches
from docx.shared import Cm
from datetime import datetime

# 创建一个新的Word文档对象
# document = Document()


# # 添加段落
# document.add_paragraph('科标业项目任务书')
# # 添加标题
# document.add_heading('This is a heading', level=1)
# 保存文档
# document.save(f'{FILE_DIR}/test_document.docx')


#!/usr/bin/python
# coding=UTF-8
# {year:'2024',}
FILE_DIR = './uploads'

#报告标题

def heading_1(document, str_b1):
    heading_1 = document.add_heading('',level=1)#返回1级标题段落对象，标题也相当于一个段落
    heading_1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 设置段落居中 
    heading_1.paragraph_format.space_before=Pt(0)#设置段前 0 磅
    heading_1.paragraph_format.space_after=Pt(0) #设置段后 0 磅
    heading_1.paragraph_format.line_spacing=1.25 #设置行间距为 1.5
    heading_1.paragraph_format.left_indent=Inches(0)#设置左缩进 1英寸
    heading_1.paragraph_format.right_indent=Inches(0)#设置右缩进 0.5 英寸
    run=heading_1.add_run(str_b1)
    run.font.name=u'宋体'    #设置为宋体
    run.font.name=u'Times New Roman'    #设置为宋体
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Times New Roman')#设置为宋体，和上边的一起使用
    run.font.size=Pt(22)#设置1级标题文字的大小为“二号” 为16磅
    run.font.color.rgb=RGBColor(0,0,0)#设置颜色为黑色

def heading_2(document, str_b2):
    heading_2 = document.add_heading('',level=2)#返回1级标题段落对象，标题也相当于一个段落
    heading_2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 设置段落居中 
    heading_2.paragraph_format.space_before=Pt(0)#设置段前 0 磅
    heading_2.paragraph_format.space_after=Pt(0) #设置段后 0 磅
    heading_2.paragraph_format.line_spacing=1.25 #设置行间距为 1.5
    heading_2.paragraph_format.left_indent=Inches(0)#设置左缩进 1英寸
    heading_2.paragraph_format.right_indent=Inches(0)#设置右缩进 0.5 英寸
    run=heading_2.add_run(str_b2)
    run.font.name=u'宋体'    #设置为宋体
    run.font.name=u'Times New Roman'    #设置为宋体
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')#设置为宋体，和上边的一起使用
    run.font.size=Pt(14)
    run.font.color.rgb=RGBColor(0,0,0)#设置颜色为黑色


def home(document, str):
    paragrapha = document.add_paragraph()
    # 首页字体设置为14磅，即四号字体
    run = paragrapha.add_run(str)
    run.font.size = Pt(14)
    # paragrapha.style.font.size = Pt(14) #设置为四号
    document.styles['Normal'].font.name = 'Times New Roman'  
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体') 
    # paragrapha.paragraph_format.first_line_indent = 2 #首行缩进2字符
    paragrapha.paragraph_format.space_before=Pt(0)#设置段前 0 磅
    paragrapha.paragraph_format.space_after=Pt(0) #设置段后 0 磅
    paragrapha.paragraph_format.line_spacing=2 #设置行间距为 2
    paragrapha.paragraph_format.left_indent = Cm(1.0) # 左缩进 1.0 厘米
    paragrapha.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  


def text(document, str):
    paragrapha = document.add_paragraph(str)
    # 将字体设置为12磅，即小四字体
    paragrapha.style.font.size = Pt(12)
    paragrapha.paragraph_format.first_line_indent = Pt(24) #首行缩进2字符 
    document.styles['Normal'].font.name = 'Times New Roman'  
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体') 
    paragrapha.paragraph_format.space_before=Pt(0)#设置段前 0 磅
    paragrapha.paragraph_format.space_after=Pt(0) #设置段后 0 磅
    paragrapha.paragraph_format.line_spacing=1.25 #设置行间距为 1.5
    # paragrapha.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐

def signature(document, str):
    paragrapha = document.add_paragraph()
    # 首页字体设置为14磅，即四号字体
    run = paragrapha.add_run(str)
    run.font.size = Pt(14)
    # paragrapha.style.font.size = Pt(14) #设置为四号
    document.styles['Normal'].font.name = 'Times New Roman'  
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体') 
    # paragrapha.paragraph_format.first_line_indent = 2 #首行缩进2字符
    paragrapha.paragraph_format.space_before=Pt(0)#设置段前 0 磅
    paragrapha.paragraph_format.space_after=Pt(0) #设置段后 0 磅
    paragrapha.paragraph_format.line_spacing=2 #设置行间距为 2
    # paragrapha.paragraph_format.left_indent = Cm(1.0) # 左缩进 1.0 厘米
    paragrapha.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  


def split_text(text):
    # 将所有的 \n\n 替换为特殊分隔符 ||
    text = text.replace('\n\n', '||')
    # 将所有的 \n 替换为空字符串
    text = text.replace('\n', '||')
    # 根据特殊分隔符 || 进行分割
    str_parts = text.split('||')
    return str_parts

def generate_docx_file(dict):

    #打开文档
    document = Document()
    document.styles['Normal'].font.name = 'Times New Roman'  
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体') 

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph = document.add_paragraph()

    str_b1 = '科标业项目任务书'
    heading_1(document, str_b1)

    #三个空行
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.25

    # paragraph = document.add_paragraph()
    # run = paragraph.add_run(u'项目类型：  ')
    # paragraph.paragraph_format.line_spacing = 2
    # paragraph.paragraph_format.left_indent = Cm(1.0)
    # paragraph.paragraph_format.space_after=Pt(0) #设置段后 0 磅
    # run.font.size=Pt(14)
    str_home = f'项目类型：  {dict["project_type"]}'
    home(document, str_home)
    str_home = f'项目编号：  {dict["project_number"]}'
    home(document, str_home)
    str_home = f'项目名称：  {dict["project_name"]}'
    home(document, str_home)
    str_home = f'项目主持部门：  {dict["host_department"]}'
    home(document, str_home)
    str_home = f'项目承担部门：  {dict["host_department"]}'
    home(document, str_home)
    str_home = f'项目负责人：  {dict["responsible_person"]}'
    home(document, str_home)


    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.25


    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'起止年限：   {dict["start_year"]}年    {dict["start_month"]} 月   ～  {dict["end_year"]}年     {dict["end_month"]}月    ')
    run.font.size=Pt(14)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after=Pt(0) #设置段后 0 磅
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 一个空行
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.25

    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'上海勘测设计研究院有限公司')
    run.bold = True
    run.font.size=Pt(14)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after=Pt(0) #设置段后 0 磅
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #添加分页
    document.add_page_break() 

    #填表说明
    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'填表说明')
    run.font.name=u'宋体'
    run.bold = True
    run.font.size=Pt(14)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after=Pt(0) #设置段后 0 磅
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 一个空行
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.25

    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'1、本任务书是公司科标业项目内部委托文件用以明确项目实施目标和项目负责人的责权利。任务书签字下发后应严肃执行。'
                            )
    run.font.size=Pt(12)
    paragraph.paragraph_format.left_indent = Pt(18)
    paragraph.paragraph_format.first_line_indent = Pt(-18)  #悬挂缩进
    paragraph.paragraph_format.space_after=Pt(0) #设置段后 0 磅
    paragraph.paragraph_format.line_spacing = 1.25


    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'2、经费预算表根据核定的费用及“公司科标业项目立项申请表”等实际情况进行填报。'
                            )
    run.font.size=Pt(12)
    paragraph.paragraph_format.left_indent = Pt(18)
    paragraph.paragraph_format.first_line_indent = Pt(-18) #悬挂缩进
    paragraph.paragraph_format.space_after=Pt(0) #设置段后 0 磅
    paragraph.paragraph_format.line_spacing = 1.25


    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'3、凡执行公司批准立项的科标业建设项目、或其主要利用公司的技术条件完成的技术成果是职务技术成果，其使用权、转让权属于公司。完成技术成果的个人有在有关技术成果文件上写明自己是技术成果完成者的权利和取得荣誉证书、奖励的权利。'
                            )
    run.font.size=Pt(12)
    paragraph.paragraph_format.left_indent = Pt(18)
    paragraph.paragraph_format.first_line_indent = Pt(-18)  #悬挂缩进
    paragraph.paragraph_format.space_after=Pt(0) #设置段后 0 磅
    paragraph.paragraph_format.line_spacing = 1.25

    #添加分页
    document.add_page_break() 


    #项目概况

    str_b2 = '一、趋势判断和需求分析'
    heading_2(document, str_b2)
    # paragraph = document.add_paragraph()
    # run = paragraph.add_run(u'一、趋势判断和需求分析')
    # run.font.name=u'宋体'
    # run.bold = True
    # run.font.size=Pt(14)
    # paragraph.paragraph_format.space_after=Pt(0) #设置段后 0 磅
    # paragraph.paragraph_format.line_spacing = 1.25

    str_text = '科技创新项目须填写国内外现状、水平和社会发展需求；科学技术价值、特色和创新点。'
    text(document, str_text)

    trend_texts = split_text(dict["trend_analysis"])
    for str_text in trend_texts:
        text(document, str_text)

    # paragraph = document.add_paragraph()
    # run = paragraph.add_run(u'科技创新项目须填写国内外现状、水平和社会发展需求；科学技术价值、特色和创新点。'
    #                         )
    # run.font.size=Pt(12)
    # paragraph.paragraph_format.line_spacing = 1.25
    # paragraph.paragraph_format.space_after=Pt(0) #设置段后 0 磅
    # paragraph.paragraph_format.first_line_indent = run.font.size * 2
    #添加分页
    document.add_page_break() 

    str_b2 = '二、项目研究内容和技术关键'
    heading_2(document, str_b2)
    str_text = '项目研究的总体目标和创新点，主要研究内容及所需要解决的技术关键、技术路线等。'
    text(document, str_text)

    research_texts = split_text(dict["research_content"])
    for str_text in research_texts:
        text(document, str_text)
    document.add_page_break() 


    str_b2 = '三、研究成果和考核指标'
    heading_2(document, str_b2)
    str_text = '''包括1.主要技术指标、形成的专利（形成不同类别专利数和可望授权专利数）、标准（标准的文件结合形成的技术标准水平）、新技术、新产品、新装置、论文专著等数量、指标和水平等；2.经济考核指标；3.人才培养情况。'''
    text(document, str_text)
    research_texts = split_text(dict["research_outcome"])
    for str_text in research_texts:
        text(document, str_text)
    document.add_page_break() 

    str_b2 = '四、年度计划和目标'
    heading_2(document, str_b2)

    str_text = '''项目的年度/季度计划及目标（按季度划分工作节点，要求明确关键的、必须实现的节点目标）'''
    text(document, str_text)

    annual_texts = split_text(dict["annual_plan"])
    for str_text in annual_texts:
        text(document, str_text)

    document.add_page_break() 

    str_b2 = '五、预期效益'
    heading_2(document, str_b2)

    str_text = '''包括直接经济效益、提高设计效率、提高公司核心竞争力、创立品牌、成果应用趋向和应用项目等）'''
    text(document, str_text)

    benefits_texts = split_text(dict["expected_benefits"])
    for str_text in benefits_texts:
        text(document, str_text)


    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.25

    str_b2 = '六、风险分析与评估'
    heading_2(document, str_b2)

    risk_texts = split_text(dict["risk_analysis"])
    for str_text in risk_texts:
        text(document, str_text)

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.25
    
    document.add_page_break() 


    #共同条款
    str_b2 = '七、共同条款'
    heading_2(document, str_b2)


    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'1.	项目负责人和项目承担部门必须每月在科研管理系统中填报项目进度及人工投入情况。若逾期不报，科技创新部有权暂停项目和终止结算。'
                            )
    run.font.size=Pt(12)
    paragraph.paragraph_format.left_indent = Cm(0.87)
    paragraph.paragraph_format.first_line_indent = Cm(-0.87)  #悬挂缩进
    paragraph.paragraph_format.space_after=Pt(0) #设置段后 0 磅
    paragraph.paragraph_format.line_spacing = 1.25


    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'2.	项目执行过程中可能影响项目顺利完成的重大事项应及时报告并按规定进行变更审批。未经变更审批流程而对项目进行了较大调整，验收时可不予通过。'
                            )
    run.font.size=Pt(12)
    paragraph.paragraph_format.left_indent = Cm(0.87)
    paragraph.paragraph_format.first_line_indent = Cm(-0.87) #悬挂缩进
    paragraph.paragraph_format.space_after=Pt(0) #设置段后 0 磅
    paragraph.paragraph_format.line_spacing = 1.25


    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'3.	项目负责人和项目承担部门因主观原因（如偏离合同内容、挪用经费、技术措施不落实等）致使计划无法执行而要求解除任务约定时，则视不同情况部分、全部或加倍退还所拨经费；科技创新部可根据调查情况提出中止合同。'
                            )
    run.font.size=Pt(12)
    paragraph.paragraph_format.left_indent = Cm(0.87)
    paragraph.paragraph_format.first_line_indent = Cm(-0.87)  #悬挂缩进
    paragraph.paragraph_format.space_after=Pt(0) #设置段后 0 磅
    paragraph.paragraph_format.line_spacing = 1.25

    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'4.	项目执行过程中，项目主持部门无故解除或不履行任务约定时，则所拨经费不得追回，并承担善后处理所发生的费用。项目主持部门提出变更任务约定有关内容时，应与项目负责人和项目承担部门协商达成书面协议后实施。'
                            )
    run.font.size=Pt(12)
    paragraph.paragraph_format.left_indent = Cm(0.87)
    paragraph.paragraph_format.first_line_indent = Cm(-0.87) #悬挂缩进
    paragraph.paragraph_format.space_after=Pt(0) #设置段后 0 磅
    paragraph.paragraph_format.line_spacing = 1.25

    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'5.	本任务书一式三份，项目主持部门执一份、项目承担部门执一份、项目负责人执一份。'
                            )
    run.font.size=Pt(12)
    paragraph.paragraph_format.left_indent = Cm(0.87)
    paragraph.paragraph_format.first_line_indent = Cm(-0.87) #悬挂缩进
    paragraph.paragraph_format.space_after=Pt(0) #设置段后 0 磅
    paragraph.paragraph_format.line_spacing = 1.25

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 2

    #任务书签约
    str_b2 = '八、任务书签约：'
    heading_2(document, str_b2)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 2

    str_home = u'项目主持部门：（公章）'
    signature(document, str_home)

    str_home = u'部门负责人：（签字）                              年      月      日'
    signature(document, str_home)

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 2

    str_home = u'项目承担部门：（公章）'
    signature(document, str_home)

    str_home = u'项目承担部门负责人：（签字）                       年      月      日'
    signature(document, str_home)

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 2

  
    str_home = u'项目负责人：（签字）                              年      月      日'
    signature(document, str_home)

    #保存文件
    # document.save('demo.docx')
    # document.save(f"{FILE_DIR}/科标业项目任务书自动化生成与润色助手{datetime.now().strftime('%Y%m%d%H%M%S')}-{dict["project_name"]}.docx")
    file_name = f"{FILE_DIR}/科标业项目任务书自动化生成与润色助手{datetime.now().strftime('%Y%m%d%H%M%S')}-{dict['project_name']}.docx"
    try:
        document.save(file_name)
        print('生成成功')
        return file_name
    except Exception as e:
        print('生成失败:', e)
        return False


if __name__ == '__main__':
    project_info = {
  'project_type': '科技创新项目',
  'project_number': 'KJ2024',
  'project_name': '基于AI图像检测算法和视觉（CV）大模型的抽水蓄能电站质量安全隐患排查和风险预警关键技术研究',
  'host_department': '三峡建工、上海院',
  'responsible_person': '张三',
  'start_year': '2024',
  'start_month': '1',
  'end_year': '2026',
  'end_month': '8',
  'trend_analysis': '随着人工智能技术的快速发展，AI图像检测算法和视觉（CV）大模型在工业安全领域的应用越来越广泛。特别是在抽水蓄能电站这样的复杂工程中，利用AI技术进行质量安全隐患排查和风险预警具有重要的现实意义和广阔的应用前景。\n\n抽水蓄能电站作为调节电力系统峰谷差的重要设施，其安全运行对保障能源供应和电网稳定至关重要。然而，由于电站建设环境复杂，施工过程中存在多种潜在的风险和隐患，如施工机械作业半径站人、卸料车后方有人员活动、高空作业保护缺失等，这些都可能导致严重的安全事故。因此，迫切需要一种高效、准确的安全监测手段来提高电站的安全运行水平。\n\n本项目旨在研究和开发一套基于AI图像检测算法和视觉（CV）大模型的抽水蓄能电站质量安全隐患排查和风险预警系统，以提高电站的安全运行水平，降低事故发生率，保障人员和设备的安全。通过创新开展以大模型作为“调度大脑”，以小模型作为“复核专家”的AI识别模式，实现对电站建设期的“少人巡检、无人巡检”模式，提升科技兴安水平。\n\n在国内外现状方面，虽然已有一些研究和应用案例，但大多数集中在通用领域，针对抽水蓄能电站这一特定领域的研究还相对较少。因此，本项目的研究具有较高的创新性和应用价值。',
  'research_content': '本项目的研究内容包括：\n\n1. 抽蓄工程风险、隐患数据集建设方法研究：将对抽蓄工程的风险和隐患进行深入分析，建立一个全面的数据集，包括图像、视频、传感器数据等，用于训练和测试AI模型。\n\n2. 视觉（CV）大模型微调方法研究：将研究如何将通用的视觉大模型微调到特定的抽蓄工程场景，以提高识别的准确性和效率。\n\n3. 视觉大、小模型调度策略研究：将研究如何有效地结合大模型和小模型，以实现更高效的风险隐患排查和预警。\n\n4. 安全风险隐患识别预警平台研究：将开发一个平台，集成上述研究成果，实现实时监控、自动报警和数据分析等功能。\n\n每个研究内容都将深入探讨，以确保项目的成功实施和技术的创新。具体研究内容如下：\n\n1. 数据集建设方法研究：数据是AI模型训练的基础，本项目将首先对抽蓄电站的安全隐患进行分类，然后收集和标注相应的图像和视频数据，构建一个多模态的数据集。此外，还将研究数据增强技术，以提高数据集的多样性和模型的泛化能力。\n\n2. 视觉（CV）大模型微调方法研究：本项目将研究如何利用现有的视觉大模型，通过迁移学习和领域适应等技术，对模型进行微调，使其更好地适应抽蓄电站的安全隐患识别任务。\n\n3. 视觉大、小模型调度策略研究：本项目将研究如何将大模型和小模型结合起来，形成一个高效的隐患排查系统。大模型负责处理常规的、通用的安全隐患识别任务，而小模型则针对特定的、复杂的安全隐患进行识别和分析。\n\n4. 安全风险隐患识别预警平台研究：本项目将开发一个集成了上述研究成果的平台，该平台能够实时接收来自电站现场的图像和传感器数据，通过AI模型进行分析，及时发现安全隐患，并进行预警。',
  'research_outcome': '本项目预期将达到以下研究成果：\n\n1. 构建一个高质量的抽蓄工程风险、隐患数据集，为AI模型的训练和测试提供丰富的数据资源。\n\n2. 开发出一套适用于抽水蓄能电站的AI图像检测算法，能够准确地识别和预警各种安全隐患。\n\n3. 形成一个有效的视觉（CV）大模型微调方法，提高模型在特定任务上的识别准确率和运算效率。\n\n4. 构建一个安全风险隐患识别预警平台，实现对电站现场的实时监控和自动预警，提高电站的安全运行水平。\n\n这些成果将为抽水蓄能电站的安全运行提供强有力的技术支持，并有望推广到其他工业安全领域。',
  'annual_plan': '本项目的年度计划如下：\n\n第一年度：\n- 第一季度：项目团队组建，项目启动会议，明确项目目标和研究内容；进行市场调研和需求分析，确定技术路线和研究方案。\n\n- 第二季度：开展数据集建设方法研究，收集和标注抽蓄电站的安全隐患数据；研究视觉（CV）大模型的微调方法，选择合适的大模型进行初步的迁移学习实验。\n\n- 第三季度：继续进行数据集建设和模型微调研究；开展视觉大、小模型调度策略研究，设计模型调度框架和策略。\n\n- 第四季度：进行安全风险隐患识别预警平台的初步开发，集成已有的AI模型和算法；在菜籽坝抽蓄电站进行小规模的现场试验和验证。\n\n第二年度：\n- 第一季度：对平台进行技术优化和系统集成，完善平台的功能和用户体验；在更多的电站场景中进行模型测试和验证。\n\n- 第二季度：对平台进行技术验证和效果评估，收集用户反馈，优化平台性能；准备项目中期报告。\n\n- 第三季度：进行项目总结和成果整理，撰写项目总结报告；准备项目验收材料。\n\n- 第四季度：成果推广和应用，将研究成果应用于其他抽水蓄能电站；撰写学术论文，申请专利和软件著作权。',
  'expected_benefits': '本项目预期将带来以下效益：\n\n1. 提高电站安全管理水平，降低事故发生率；通过实时监控和自动预警，及时发现和处理安全隐患，减少事故发生的可能性。\n\n2. 减少人力成本，提高隐患排查和风险预警的效率；自动化的隐患排查系统将减少对人工巡检的依赖，降低人力成本。\n\n3. 通过技术创新，提升集团在抽水蓄能电站领域的竞争力；本项目的研究成果将增强集团在新能源领域的技术实力和市场竞争力。\n\n4. 为集团在新能源领域的可持续发展提供技术支持；本项目的技术成果将有助于集团在新能源领域的长期发展。\n\n此外，项目还将产生一系列学术论文、专利和软件著作权，增强集团的知识产权储备。',
  'risk_analysis': '本项目可能面临的风险包括：\n\n1. 技术风险：新技术的研发可能遇到预期之外的难题；AI模型的训练和优化是一个复杂的过程，可能会遇到数据质量、算法效率等问题。\n\n2. 市场风险：市场需求的变化可能影响项目的推广和应用；如果市场对AI在工业安全领域的应用接受度不高，可能会影响项目的推广。\n\n3. 财务风险：项目成本的超支可能影响项目的经济效益；项目实施过程中可能会遇到预算外的支出，导致成本超支。\n\n4. 法律风险：知识产权保护和合同纠纷可能对项目造成不利影响；项目涉及的技术和产品需要严格的知识产权保护，否则可能会面临侵权风险。\n\n针对这些风险，项目团队将采取相应的风险管理和应对措施，确保项目的顺利进行。具体措施包括：\n\n1. 技术风险管理：通过与高校和研究机构合作，引入专家咨询，确保技术难题得到及时解决；同时，建立技术储备和替代方案，以应对可能的技术挑战。\n\n2. 市场风险管理：通过市场调研，及时调整项目方向和产品功能，以适应市场需求；同时，加强市场推广和用户培训，提高市场接受度。\n\n3. 财务风险管理：通过严格的预算控制和成本管理，确保项目成本不超支；同时，寻求政府补贴和行业资助，降低财务压力。\n\n4. 法律风险管理：通过加强知识产权保护和合同管理，避免法律纠纷；同时，建立法律顾问团队，处理可能的法律问题。'
}
    generate_docx_file(project_info)
